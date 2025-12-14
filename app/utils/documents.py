"""
Document parsing utilities for PDF and DOCX files.

Provides extraction of text, metadata, and images from documents.
"""

import io
import os
import base64
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from urllib.parse import urlparse

import httpx

from app.utils.logger import get_logger
from app.utils.url_validator import validate_url

logger = get_logger(__name__)

# Document extensions we can handle
DOCUMENT_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword',
}

# Content types we can handle
DOCUMENT_CONTENT_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
}


class DocumentParseError(Exception):
    """Raised when document parsing fails."""
    pass


def is_document_url(url: str, content_type: Optional[str] = None) -> Tuple[bool, Optional[str]]:
    """
    Check if a URL points to a document we can parse.

    Args:
        url: URL to check
        content_type: Optional content-type header from response

    Returns:
        Tuple of (is_document, document_type)
    """
    # Check by content type first (most reliable)
    if content_type:
        content_type_lower = content_type.lower().split(';')[0].strip()
        if content_type_lower in DOCUMENT_CONTENT_TYPES:
            return True, DOCUMENT_CONTENT_TYPES[content_type_lower]

    # Fall back to extension check
    parsed = urlparse(url)
    path = parsed.path.lower()

    for ext, _ in DOCUMENT_EXTENSIONS.items():
        if path.endswith(ext):
            doc_type = 'pdf' if ext == '.pdf' else 'docx' if ext == '.docx' else 'doc'
            return True, doc_type

    return False, None


async def download_document(url: str, timeout: int = 30000) -> Tuple[bytes, str]:
    """
    Download a document from URL.

    Args:
        url: Document URL
        timeout: Timeout in milliseconds

    Returns:
        Tuple of (document_bytes, content_type)
    """
    # Validate URL first
    is_valid, error = validate_url(url)
    if not is_valid:
        raise DocumentParseError(f"URL blocked: {error}")

    timeout_seconds = timeout / 1000

    async with httpx.AsyncClient(timeout=timeout_seconds, follow_redirects=True) as client:
        response = await client.get(url)
        response.raise_for_status()

        content_type = response.headers.get('content-type', '')
        return response.content, content_type


def parse_pdf(content: bytes, extract_images: bool = False) -> Dict[str, Any]:
    """
    Parse a PDF document and extract text and metadata.

    Args:
        content: PDF file bytes
        extract_images: Whether to extract embedded images

    Returns:
        Dictionary with text, metadata, and optionally images
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise DocumentParseError("pymupdf not installed. Install with: pip install pymupdf")

    result = {
        "text": "",
        "markdown": "",
        "metadata": {},
        "page_count": 0,
        "images": []
    }

    try:
        # Open PDF from bytes
        doc = fitz.open(stream=content, filetype="pdf")

        result["page_count"] = len(doc)

        # Extract metadata
        metadata = doc.metadata
        if metadata:
            result["metadata"] = {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "keywords": metadata.get("keywords", ""),
                "creator": metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
                "creationDate": metadata.get("creationDate", ""),
                "modDate": metadata.get("modDate", ""),
            }

        # Extract text from each page
        text_parts = []
        markdown_parts = []

        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            text_parts.append(page_text)

            # Create markdown with page breaks
            if page_num > 0:
                markdown_parts.append("\n\n---\n\n")
            markdown_parts.append(f"## Page {page_num + 1}\n\n{page_text}")

            # Extract images if requested
            if extract_images:
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]

                            result["images"].append({
                                "page": page_num + 1,
                                "index": img_index,
                                "format": image_ext,
                                "data": base64.b64encode(image_bytes).decode(),
                                "width": base_image.get("width"),
                                "height": base_image.get("height")
                            })
                    except Exception as e:
                        logger.warning("pdf_image_extraction_failed",
                                     page=page_num, index=img_index, error=str(e))

        result["text"] = "\n\n".join(text_parts)
        result["markdown"] = "".join(markdown_parts)

        doc.close()

        logger.info("pdf_parsed", page_count=result["page_count"],
                   text_length=len(result["text"]), image_count=len(result["images"]))

        return result

    except Exception as e:
        logger.error("pdf_parse_failed", error=str(e))
        raise DocumentParseError(f"Failed to parse PDF: {str(e)}")


def parse_docx(content: bytes, extract_images: bool = False) -> Dict[str, Any]:
    """
    Parse a DOCX document and extract text and metadata.

    Args:
        content: DOCX file bytes
        extract_images: Whether to extract embedded images

    Returns:
        Dictionary with text, metadata, and optionally images
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        raise DocumentParseError("python-docx not installed. Install with: pip install python-docx")

    result = {
        "text": "",
        "markdown": "",
        "metadata": {},
        "paragraph_count": 0,
        "table_count": 0,
        "images": []
    }

    try:
        # Open DOCX from bytes
        doc = Document(io.BytesIO(content))

        # Extract core properties (metadata)
        core_props = doc.core_properties
        result["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "category": core_props.category or "",
            "comments": core_props.comments or "",
        }

        # Extract paragraphs
        text_parts = []
        markdown_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

                # Convert to markdown based on style
                style_name = para.style.name.lower() if para.style else ""

                if "heading 1" in style_name:
                    markdown_parts.append(f"# {text}")
                elif "heading 2" in style_name:
                    markdown_parts.append(f"## {text}")
                elif "heading 3" in style_name:
                    markdown_parts.append(f"### {text}")
                elif "heading" in style_name:
                    markdown_parts.append(f"#### {text}")
                elif "list" in style_name or "bullet" in style_name:
                    markdown_parts.append(f"- {text}")
                else:
                    markdown_parts.append(text)

        result["paragraph_count"] = len(doc.paragraphs)

        # Extract tables
        for table_num, table in enumerate(doc.tables):
            markdown_parts.append(f"\n**Table {table_num + 1}:**\n")

            table_md = []
            for row_num, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                table_md.append("| " + " | ".join(cells) + " |")

                # Add header separator after first row
                if row_num == 0:
                    table_md.append("|" + "|".join(["---"] * len(cells)) + "|")

            markdown_parts.append("\n".join(table_md))

        result["table_count"] = len(doc.tables)

        # Extract images if requested
        if extract_images:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob

                        # Determine image format from content type
                        content_type = image_part.content_type
                        ext = content_type.split('/')[-1] if '/' in content_type else 'png'

                        result["images"].append({
                            "format": ext,
                            "data": base64.b64encode(image_bytes).decode(),
                            "content_type": content_type
                        })
                    except Exception as e:
                        logger.warning("docx_image_extraction_failed", error=str(e))

        result["text"] = "\n\n".join(text_parts)
        result["markdown"] = "\n\n".join(markdown_parts)

        logger.info("docx_parsed", paragraph_count=result["paragraph_count"],
                   table_count=result["table_count"], text_length=len(result["text"]),
                   image_count=len(result["images"]))

        return result

    except Exception as e:
        logger.error("docx_parse_failed", error=str(e))
        raise DocumentParseError(f"Failed to parse DOCX: {str(e)}")


async def parse_document_url(
    url: str,
    formats: List[str],
    timeout: int = 30000
) -> Dict[str, Any]:
    """
    Download and parse a document from URL.

    Args:
        url: Document URL
        formats: Requested output formats (markdown, text, metadata, media)
        timeout: Timeout in milliseconds

    Returns:
        Dictionary with parsed content
    """
    logger.info("document_parse_started", url=url, formats=formats)

    # Download the document
    content, content_type = await download_document(url, timeout)

    # Determine document type
    is_doc, doc_type = is_document_url(url, content_type)

    if not is_doc:
        raise DocumentParseError(f"Unsupported document type: {content_type}")

    # Parse based on document type
    extract_images = "media" in formats or "images" in formats

    if doc_type == "pdf":
        parsed = parse_pdf(content, extract_images)
    elif doc_type in ("docx", "doc"):
        if doc_type == "doc":
            logger.warning("doc_format_limited",
                         message="Legacy .doc format has limited support, consider converting to .docx")
        parsed = parse_docx(content, extract_images)
    else:
        raise DocumentParseError(f"Unsupported document type: {doc_type}")

    # Build result based on requested formats
    result = {
        "document_type": doc_type,
    }

    if "markdown" in formats:
        result["markdown"] = parsed.get("markdown", "")
        result["quality_score"] = _calculate_document_quality(parsed)
        result["extraction_method"] = f"{doc_type}_parser"

    if "text" in formats or "html" in formats:
        result["text"] = parsed.get("text", "")

    if "metadata" in formats:
        metadata = parsed.get("metadata", {})
        metadata["sourceURL"] = url
        metadata["documentType"] = doc_type
        metadata["pageCount"] = parsed.get("page_count")
        metadata["paragraphCount"] = parsed.get("paragraph_count")
        metadata["tableCount"] = parsed.get("table_count")
        result["metadata"] = metadata

    if "media" in formats or "images" in formats:
        result["images"] = parsed.get("images", [])

    logger.info("document_parse_completed", url=url, doc_type=doc_type)

    return result


def _calculate_document_quality(parsed: Dict[str, Any]) -> float:
    """
    Calculate a quality score for parsed document content.

    Args:
        parsed: Parsed document data

    Returns:
        Quality score from 0.0 to 1.0
    """
    text = parsed.get("text", "")
    markdown = parsed.get("markdown", "")

    if not text:
        return 0.0

    score = 0.0

    # Length score (0-0.3)
    text_length = len(text)
    if text_length > 10000:
        score += 0.3
    elif text_length > 5000:
        score += 0.25
    elif text_length > 1000:
        score += 0.2
    elif text_length > 500:
        score += 0.15
    elif text_length > 100:
        score += 0.1
    else:
        score += 0.05

    # Structure score (0-0.3)
    has_headings = "#" in markdown
    has_lists = "- " in markdown or "* " in markdown
    has_tables = "|" in markdown and "---" in markdown

    if has_headings:
        score += 0.1
    if has_lists:
        score += 0.1
    if has_tables:
        score += 0.1

    # Metadata score (0-0.2)
    metadata = parsed.get("metadata", {})
    if metadata.get("title"):
        score += 0.1
    if metadata.get("author"):
        score += 0.05
    if metadata.get("subject") or metadata.get("keywords"):
        score += 0.05

    # Content quality score (0-0.2)
    # Check for reasonable word density
    words = text.split()
    word_count = len(words)

    if word_count > 500:
        score += 0.1
    elif word_count > 100:
        score += 0.05

    # Check for non-garbage content (average word length)
    if words:
        avg_word_length = sum(len(w) for w in words) / len(words)
        if 3 < avg_word_length < 12:
            score += 0.1

    return min(score, 1.0)
